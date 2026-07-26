#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import io
import json
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw
from lxml import etree

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "tests" / "fixtures" / "docx"
FIXED_DT = (2026, 7, 20, 0, 0, 0)
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
CT = "http://schemas.openxmlformats.org/package/2006/content-types"


def set_font(run, size=10, bold=False):
    run.font.name = "Liberation Sans"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Sans")
    run.font.size = Pt(size)
    run.bold = bold


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, 14, True)


def base_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Liberation Sans"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Liberation Sans")
    style.font.size = Pt(10)
    return doc


def make_logo(path: Path):
    image = Image.new("RGB", (360, 100), "white")
    d = ImageDraw.Draw(image)
    d.rectangle((2, 2, 357, 97), outline="black", width=3)
    d.text((20, 35), "DOKKOMPLEKT GOLDEN", fill="black")
    image.save(path)


def merged_tables_sections(path: Path):
    doc = base_doc()
    add_title(doc, "GOLDEN: объединённые таблицы и секции")
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    merged = table.cell(0, 0).merge(table.cell(0, 3))
    merged.text = "{{organization.name}} — сводный документ"
    for c in table.rows[0].cells:
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    headers = ["Наименование", "Количество", "Цена", "Сумма"]
    for i, value in enumerate(headers):
        table.cell(1, i).text = value
    for row in range(2, 5):
        table.cell(row, 0).text = f"Позиция {row-1}"
        table.cell(row, 1).text = "{{#each items}}{{quantity}}{{/each}}"
        table.cell(row, 2).text = "{{price}}"
        table.cell(row, 3).text = "{{sum(items.amount)}}"
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, 9)
    doc.add_paragraph("Условный блок: {{#if contract.number}}Договор № {{contract.number}}{{/if}}")
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.header.paragraphs[0].text = "Колонтитул: {{organization.name}}"
    section.footer.paragraphs[0].text = "Страница поля PAGE"
    p = doc.add_paragraph("Альбомная секция с длинной строкой и кириллицей: ")
    set_font(p.add_run("ФИО, ИНН, КПП, СНИЛС, кадастровый номер, VIN."), 10, True)
    wide = doc.add_table(rows=3, cols=8)
    wide.style = "Table Grid"
    for r in range(3):
        for c in range(8):
            wide.cell(r, c).text = f"R{r+1}C{c+1}"
    doc.save(path)


def long_pagination(path: Path):
    doc = base_doc()
    add_title(doc, "GOLDEN: многостраничная таблица")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["№", "Дата", "Описание", "Ответственный", "Статус"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for n in range(1, 86):
        cells = table.add_row().cells
        values = [str(n), f"{(n%28)+1:02}.07.2026", f"Строка {n}: проверка переносов и пагинации документа.", "{{person.full_name}}", "Готово"]
        for c, value in zip(cells, values):
            c.text = value
    doc.save(path)


def image_header_footer(path: Path, image_path: Path):
    doc = base_doc()
    section = doc.sections[0]
    section.header.paragraphs[0].add_run().add_picture(str(image_path), width=Cm(6.5))
    section.footer.paragraphs[0].text = "Конфиденциально — {{document.number}}"
    add_title(doc, "GOLDEN: изображения в теле и колонтитулах")
    p = doc.add_paragraph("Печать организации: ")
    p.add_run().add_picture(str(image_path), width=Cm(5.0))
    doc.add_paragraph("Маркер для замены: {{image org.stamp}}")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Подпись"
    table.cell(0, 1).text = "{{image person.signature}}"
    table.cell(1, 0).text = "Факсимиле"
    table.cell(1, 1).text = "{{image org.facsimile}}"
    doc.save(path)


def formula_conditions(path: Path):
    doc = base_doc()
    add_title(doc, "GOLDEN: формулы, условия и буквальные фигурные скобки")
    values = [
        "Итого: {{sum(items.amount)}}",
        "Количество: {{count(items)}}",
        "Рабочих дней: {{working_days(document.start_date, document.end_date)}}",
        "{{#unless subject.address}}Адрес не указан{{else}}{{subject.address}}{{/unless}}",
        r"Литеральный код: \{{this_is_not_a_placeholder\}}",
        "Формула: {{(price * quantity) + tax}}",
        "Текст с << и >> не является плейсхолдером.",
    ]
    for value in values:
        p = doc.add_paragraph(value)
        for run in p.runs:
            set_font(run, 10)
    doc.save(path)


def content_controls_fields(path: Path):
    doc = base_doc()
    add_title(doc, "GOLDEN: content controls и поля Word")
    doc.add_paragraph("PLACEHOLDER_SDT")
    doc.add_paragraph("PLACEHOLDER_FIELD")
    doc.add_paragraph("PLACEHOLDER_TRACKED")
    doc.save(path)
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with zipfile.ZipFile(path) as zin:
            zin.extractall(td)
        document_path = td / "word/document.xml"
        tree = etree.parse(str(document_path))
        ns = {"w": W}
        for paragraph in tree.xpath("//w:p", namespaces=ns):
            text = "".join(paragraph.xpath(".//w:t/text()", namespaces=ns))
            if text == "PLACEHOLDER_SDT":
                paragraph.clear()
                sdt = etree.SubElement(paragraph, f"{{{W}}}sdt")
                props = etree.SubElement(sdt, f"{{{W}}}sdtPr")
                alias = etree.SubElement(props, f"{{{W}}}alias")
                alias.set(f"{{{W}}}val", "Document Number")
                tag = etree.SubElement(props, f"{{{W}}}tag")
                tag.set(f"{{{W}}}val", "document.number")
                content = etree.SubElement(sdt, f"{{{W}}}sdtContent")
                run = etree.SubElement(content, f"{{{W}}}r")
                t = etree.SubElement(run, f"{{{W}}}t")
                t.text = "{{document.number}}"
            elif text == "PLACEHOLDER_FIELD":
                paragraph.clear()
                fld = etree.SubElement(paragraph, f"{{{W}}}fldSimple")
                fld.set(f"{{{W}}}instr", 'DATE \\@ "dd.MM.yyyy"')
                run = etree.SubElement(fld, f"{{{W}}}r")
                t = etree.SubElement(run, f"{{{W}}}t")
                t.text = "20.07.2026"
            elif text == "PLACEHOLDER_TRACKED":
                paragraph.clear()
                deleted = etree.SubElement(paragraph, f"{{{W}}}del")
                deleted.set(f"{{{W}}}author", "QA")
                deleted.set(f"{{{W}}}date", "2026-07-20T00:00:00Z")
                dr = etree.SubElement(deleted, f"{{{W}}}r")
                dt = etree.SubElement(dr, f"{{{W}}}delText")
                dt.text = "Старое значение"
                inserted = etree.SubElement(paragraph, f"{{{W}}}ins")
                inserted.set(f"{{{W}}}author", "QA")
                inserted.set(f"{{{W}}}date", "2026-07-20T00:00:00Z")
                ir = etree.SubElement(inserted, f"{{{W}}}r")
                it = etree.SubElement(ir, f"{{{W}}}t")
                it.text = "Новое значение"
        tree.write(str(document_path), xml_declaration=True, encoding="UTF-8", standalone="yes")
        repack(td, path)


def comments_and_links(path: Path):
    doc = base_doc()
    add_title(doc, "GOLDEN: комментарии и внешние связи")
    doc.add_paragraph("COMMENT_ANCHOR")
    doc.add_paragraph("Документ должен сохранять комментарии и связи после рендера.")
    doc.save(path)
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with zipfile.ZipFile(path) as zin:
            zin.extractall(td)
        document_path = td / "word/document.xml"
        tree = etree.parse(str(document_path))
        ns = {"w": W}
        for p in tree.xpath("//w:p", namespaces=ns):
            if "".join(p.xpath(".//w:t/text()", namespaces=ns)) == "COMMENT_ANCHOR":
                p.clear()
                start = etree.SubElement(p, f"{{{W}}}commentRangeStart")
                start.set(f"{{{W}}}id", "0")
                r = etree.SubElement(p, f"{{{W}}}r")
                t = etree.SubElement(r, f"{{{W}}}t")
                t.text = "Проверяемый фрагмент"
                end = etree.SubElement(p, f"{{{W}}}commentRangeEnd")
                end.set(f"{{{W}}}id", "0")
                rr = etree.SubElement(p, f"{{{W}}}r")
                ref = etree.SubElement(rr, f"{{{W}}}commentReference")
                ref.set(f"{{{W}}}id", "0")
        tree.write(str(document_path), xml_declaration=True, encoding="UTF-8", standalone="yes")
        comments = etree.Element(f"{{{W}}}comments", nsmap={"w": W})
        comment = etree.SubElement(comments, f"{{{W}}}comment")
        comment.set(f"{{{W}}}id", "0")
        comment.set(f"{{{W}}}author", "Dokkomplekt QA")
        comment.set(f"{{{W}}}date", "2026-07-20T00:00:00Z")
        cp = etree.SubElement(comment, f"{{{W}}}p")
        cr = etree.SubElement(cp, f"{{{W}}}r")
        ct = etree.SubElement(cr, f"{{{W}}}t")
        ct.text = "Комментарий должен сохраниться структурно."
        (td / "word/comments.xml").write_bytes(etree.tostring(comments, xml_declaration=True, encoding="UTF-8", standalone=True))
        rels_path = td / "word/_rels/document.xml.rels"
        rels = etree.parse(str(rels_path))
        rel = etree.SubElement(rels.getroot(), f"{{{PKG_REL}}}Relationship")
        rel.set("Id", "rIdComments")
        rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
        rel.set("Target", "comments.xml")
        rels.write(str(rels_path), xml_declaration=True, encoding="UTF-8", standalone="yes")
        ct_path = td / "[Content_Types].xml"
        ctt = etree.parse(str(ct_path))
        override = etree.SubElement(ctt.getroot(), f"{{{CT}}}Override")
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
        ctt.write(str(ct_path), xml_declaration=True, encoding="UTF-8", standalone="yes")
        repack(td, path)


def repack(folder: Path, target: Path):
    temp = target.with_suffix(".tmp.docx")
    with zipfile.ZipFile(temp, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
        for file in sorted(p for p in folder.rglob("*") if p.is_file()):
            rel = file.relative_to(folder).as_posix()
            info = zipfile.ZipInfo(rel, FIXED_DT)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o100644 << 16
            zout.writestr(info, file.read_bytes())
    temp.replace(target)


def normalize_docx(path: Path):
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        with zipfile.ZipFile(path) as zin:
            zin.extractall(td)
        # Remove volatile core properties while preserving valid metadata.
        core = td / "docProps/core.xml"
        if core.exists():
            tree = etree.parse(str(core))
            ns = {
                "dc": "http://purl.org/dc/elements/1.1/",
                "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
                "dcterms": "http://purl.org/dc/terms/",
            }
            for node in tree.xpath("//*[local-name()='created' or local-name()='modified' or local-name()='lastModifiedBy' or local-name()='revision']"):
                node.getparent().remove(node)
            creator = tree.find("dc:creator", ns)
            if creator is not None:
                creator.text = "Dokkomplekt QA"
            tree.write(str(core), xml_declaration=True, encoding="UTF-8", standalone="yes")
        repack(td, path)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    logo = OUT / "_golden_logo.png"
    make_logo(logo)
    builders = {
        "merged_tables_sections.docx": lambda p: merged_tables_sections(p),
        "long_pagination.docx": lambda p: long_pagination(p),
        "images_headers_footers.docx": lambda p: image_header_footer(p, logo),
        "formula_conditions.docx": lambda p: formula_conditions(p),
        "content_controls_fields_tracked.docx": lambda p: content_controls_fields(p),
        "comments_relationships.docx": lambda p: comments_and_links(p),
    }
    for name, build in builders.items():
        path = OUT / name
        build(path)
        normalize_docx(path)
    logo.unlink(missing_ok=True)
    manifest = {
        "schema": 1,
        "generated_by": "scripts/generate_docx_golden_corpus.py",
        "fixtures": {
            path.name: {
                "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
                "size_bytes": path.stat().st_size,
            }
            for path in sorted(OUT.glob("*.docx"))
        },
    }
    (OUT / "corpus-manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
