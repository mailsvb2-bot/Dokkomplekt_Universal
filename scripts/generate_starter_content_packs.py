#!/usr/bin/env python3
"""Generate deterministic, visibly marked starter DOCX templates for optional packs.

These templates are intentionally draft-only. They provide a working layout and
canonical placeholders without pretending to be legally approved forms.
"""
from __future__ import annotations

import datetime as dt
import hashlib
import json
import re
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

ROOT = Path(__file__).resolve().parents[1]
PACK_ROOT = ROOT / "content-packs"
FIXED_ZIP_TIME = (1980, 1, 1, 0, 0, 0)
FIELD_TOKEN_RE = re.compile(r"[a-z][a-z0-9_]*(?:\.[a-z][a-z0-9_]*)+")

DISCLAIMER = (
    "STARTER-КАРКАС. НЕ ЯВЛЯЕТСЯ УТВЕРЖДЁННОЙ, НОРМАТИВНОЙ ИЛИ ЮРИДИЧЕСКИ "
    "ЗНАЧИМОЙ ФОРМОЙ. Перед использованием документ должен быть проверен и "
    "утверждён уполномоченным специалистом вашей организации."
)


@dataclass(frozen=True)
class TemplateSpec:
    pack: str
    filename: str
    title: str
    subtitle: str
    paragraphs: tuple[str, ...]
    table: tuple[tuple[str, str], ...] = ()
    signatures: tuple[str, ...] = ()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def configure_document(doc: Document) -> None:
    fixed = dt.datetime(2026, 1, 1, 0, 0, 0)
    props = doc.core_properties
    props.author = "Dokkomplekt starter-pack generator"
    props.last_modified_by = "Dokkomplekt starter-pack generator"
    props.created = fixed
    props.modified = fixed
    props.revision = 1
    section = doc.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(18)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    for style_name in ("Title", "Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Arial"
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DISCLAIMER)
    run.bold = True
    run.font.size = Pt(8)


def add_title(doc: Document, spec: TemplateSpec) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(spec.title)
    run.bold = True
    run.font.size = Pt(15)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(spec.subtitle)
    r2.italic = True
    r2.font.size = Pt(9.5)


def add_key_value_table(doc: Document, rows: tuple[tuple[str, str], ...]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row().cells
        row[0].width = Mm(55)
        row[1].width = Mm(115)
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row[0], "EDEDED")
        p0 = row[0].paragraphs[0]
        p0.add_run(label).bold = True
        row[1].paragraphs[0].add_run(value)


def add_signatures(doc: Document, signatures: tuple[str, ...]) -> None:
    if not signatures:
        return
    doc.add_paragraph()
    for label in signatures:
        p = doc.add_paragraph()
        p.add_run(f"{label}: __________________ / __________________")


def save_deterministic(doc: Document, target: Path) -> None:
    target.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp:
        raw = Path(tmp) / "raw.docx"
        normalized = Path(tmp) / "normalized.docx"
        doc.save(raw)
        with zipfile.ZipFile(raw, "r") as src, zipfile.ZipFile(
            normalized, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
        ) as dst:
            for name in sorted(src.namelist()):
                data = src.read(name)
                info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o100644 << 16
                dst.writestr(info, data)
        shutil.copyfile(normalized, target)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def build(spec: TemplateSpec) -> Path:
    doc = Document()
    configure_document(doc)
    add_title(doc, spec)
    add_key_value_table(doc, spec.table)
    for text in spec.paragraphs:
        doc.add_paragraph(text)
    add_signatures(doc, spec.signatures)
    target = PACK_ROOT / spec.pack / "templates" / spec.filename
    save_deterministic(doc, target)
    return target


SPECS = (
    TemplateSpec(
        "tier1-hr-ru", "employment_contract.docx", "ТРУДОВОЙ ДОГОВОР № {{employee.contract_number}}",
        "Черновой starter-шаблон для настройки собственного утверждённого документа",
        (
            "{{org.name}}, именуемое далее «Работодатель», и {{employee.name}}, именуемый(ая) далее «Работник», фиксируют основные сведения для последующей юридической редакции.",
            "Должность: {{employee.position}}. Подразделение: {{employee.department}}.",
            "Дата начала работы: {{employee.hire_date}}. Оклад: {{employee.salary}}.",
            "Условия труда, режим, гарантии, компенсации, ответственность и порядок прекращения договора должны быть внесены из утверждённой формы работодателя.",
        ),
        (("Дата документа", "{{document.date}}"), ("Организация", "{{org.name}}"), ("Сотрудник", "{{employee.name}}")),
        ("Работодатель", "Работник"),
    ),
    TemplateSpec(
        "tier1-hr-ru", "employment_order.docx", "ПРИКАЗ О ПРИЁМЕ НА РАБОТУ № {{hr.order_number}}",
        "Черновой starter-шаблон",
        (
            "Принять {{employee.name}} на должность {{employee.position}} в подразделение {{employee.department}} с {{employee.hire_date}}.",
            "Основание: трудовой договор № {{employee.contract_number}}.",
            "Поля, формулировки основания и реквизиты утверждаются работодателем до публикации шаблона.",
        ),
        (("Дата приказа", "{{hr.order_date}}"), ("Табельный номер", "{{employee.tab_number}}")),
        ("Руководитель", "С приказом ознакомлен(а)"),
    ),
    TemplateSpec(
        "tier1-hr-ru", "personal_data_consent.docx", "СОГЛАСИЕ НА ОБРАБОТКУ ПЕРСОНАЛЬНЫХ ДАННЫХ",
        "Черновой starter-шаблон — требуется проверка ответственного за ПДн",
        (
            "Субъект персональных данных: {{employee.name}}.",
            "Оператор: {{org.name}}, ИНН {{org.inn}}, адрес {{org.legal_address}}.",
            "Перечень данных, цели, способы обработки, сроки хранения, порядок отзыва и трансграничная передача должны быть определены локальными актами оператора.",
        ),
        (("Дата", "{{document.date}}"),),
        ("Субъект персональных данных",),
    ),
    TemplateSpec(
        "tier1-hr-ru", "familiarization_sheet.docx", "ЛИСТ ОЗНАКОМЛЕНИЯ",
        "Черновой starter-шаблон",
        (
            "Сотрудник {{employee.name}}, должность {{employee.position}}, подтверждает ознакомление с перечнем локальных актов, утверждённым организацией.",
            "Перечень актов и порядок ознакомления должны быть заполнены работодателем перед публикацией шаблона.",
        ),
        (("Организация", "{{org.name}}"), ("Дата", "{{document.date}}")),
        ("Сотрудник", "Ответственный"),
    ),
    TemplateSpec(
        "tier1-legal-ru", "contract.docx", "ДОГОВОР № {{contract.number}}",
        "Черновой starter-шаблон — не заменяет юридическую экспертизу",
        (
            "{{org.name}} и {{counterparty.name}} фиксируют реквизиты будущего договора.",
            "Предмет: {{contract.subject}}.",
            "Цена: {{contract.amount}} {{contract.currency}}.",
            "Срок: с {{contract.start_date}} по {{contract.end_date}}.",
            "Права, обязанности, ответственность, заверения, порядок приёмки, расторжения и разрешения споров должны быть внесены юристом из утверждённой редакции.",
        ),
        (("Дата", "{{contract.date}}"), ("Сторона 1", "{{org.name}}"), ("Сторона 2", "{{counterparty.name}}")),
        ("Сторона 1", "Сторона 2"),
    ),
    TemplateSpec(
        "tier1-legal-ru", "acceptance_act.docx", "АКТ № {{document.number}}",
        "Черновой starter-шаблон к договору",
        (
            "Стороны {{org.name}} и {{counterparty.name}} составили настоящий акт к договору № {{contract.number}} от {{contract.date}}.",
            "Описание принятого результата: {{contract.subject}}.",
            "Стоимость по акту: {{contract.amount}} {{contract.currency}}.",
            "Оговорки, объём, сроки и порядок приёмки должны быть согласованы в утверждённом шаблоне.",
        ),
        (("Дата акта", "{{document.date}}"),),
        ("Сторона 1", "Сторона 2"),
    ),
    TemplateSpec(
        "tier1-legal-ru", "claim.docx", "ПРЕТЕНЗИЯ № {{document.number}}",
        "Черновой starter-шаблон — правовая позиция требует проверки юристом",
        (
            "Отправитель: {{org.name}}. Получатель: {{counterparty.name}}.",
            "Основание: договор № {{contract.number}} от {{contract.date}}.",
            "Предмет требования: {{legal.claim_subject}}.",
            "Сумма требования: {{legal.claim_amount}}.",
            "Нормативное обоснование, срок ответа, расчёт и приложения должны быть заполнены юристом.",
        ),
        (("Дата", "{{document.date}}"),),
        ("Уполномоченное лицо",),
    ),
    TemplateSpec(
        "tier1-legal-ru", "cover_letter.docx", "СОПРОВОДИТЕЛЬНОЕ ПИСЬМО № {{document.number}}",
        "Черновой starter-шаблон",
        (
            "{{org.name}} направляет в адрес {{counterparty.name}} комплект документов по вопросу: {{contract.subject}}.",
            "Перечень приложений и способ доставки должны быть заполнены перед отправкой.",
        ),
        (("Дата", "{{document.date}}"),),
        ("Отправитель",),
    ),
    TemplateSpec(
        "tier1-accounting-ru", "invoice.docx", "СЧЁТ № {{accounting.invoice_number}}",
        "Черновой starter-шаблон — не является первичным учётным документом",
        (
            "Поставщик: {{org.name}}, ИНН {{org.inn}}, КПП {{org.kpp}}.",
            "Покупатель: {{counterparty.name}}.",
            "Назначение платежа и номенклатура должны быть заполнены из реестра или утверждённого шаблона.",
        ),
        (("Дата", "{{accounting.invoice_date}}"), ("Итого", "{{amount.total}} {{amount.currency}}"), ("НДС", "{{amount.vat}}")),
        ("Ответственный",),
    ),
    TemplateSpec(
        "tier1-accounting-ru", "service_act.docx", "АКТ ОКАЗАННЫХ УСЛУГ № {{document.number}}",
        "Черновой starter-шаблон",
        (
            "Исполнитель: {{org.name}}. Заказчик: {{counterparty.name}}.",
            "Основание: договор № {{contract.number}} от {{contract.date}}.",
            "Содержание услуг: {{contract.subject}}.",
            "Стоимость: {{amount.total}} {{amount.currency}}, НДС: {{amount.vat}}.",
            "Объём, период и условия приёмки должны быть подтверждены первичными данными.",
        ),
        (("Дата", "{{document.date}}"),),
        ("Исполнитель", "Заказчик"),
    ),
    TemplateSpec(
        "tier1-accounting-ru", "reconciliation.docx", "АКТ СВЕРКИ ВЗАИМНЫХ РАСЧЁТОВ",
        "Черновой starter-шаблон — итоговые суммы должны рассчитываться из реестра",
        (
            "Стороны: {{org.name}} и {{counterparty.name}}.",
            "Период сверки и строки операций должны быть сформированы из XLSX/CSV-реестра организации.",
            "Итоговое сальдо: {{amount.total}} {{amount.currency}}.",
        ),
        (("Дата формирования", "{{document.date}}"),),
        ("Сторона 1", "Сторона 2"),
    ),
)



def fields_in_spec(spec: TemplateSpec) -> list[str]:
    text = "\n".join(
        [spec.title, spec.subtitle, *spec.paragraphs]
        + [item for row in spec.table for item in row]
        + list(spec.signatures)
    )
    return sorted(set(FIELD_TOKEN_RE.findall(text.lower())))

def update_manifest(pack: str, file_hashes: dict[str, str]) -> None:
    manifest_path = PACK_ROOT / pack / "pack.json"
    data = json.loads(manifest_path.read_text("utf-8"))
    data["version"] = "0.3.0"
    data["status"] = "starter"
    data["usage_mode"] = "draft_only"
    data["requires_organization_review"] = True
    for slot in data["template_slots"]:
        filename = {
            "hr.employment_contract": "employment_contract.docx",
            "hr.employment_order": "employment_order.docx",
            "hr.personal_data_consent": "personal_data_consent.docx",
            "hr.familiarization_sheet": "familiarization_sheet.docx",
            "legal.contract": "contract.docx",
            "legal.acceptance_act": "acceptance_act.docx",
            "legal.claim": "claim.docx",
            "legal.cover_letter": "cover_letter.docx",
            "accounting.invoice": "invoice.docx",
            "accounting.service_act": "service_act.docx",
            "accounting.reconciliation": "reconciliation.docx",
        }[slot["document_id"]]
        spec = next(item for item in SPECS if item.pack == pack and item.filename == filename)
        references = fields_in_spec(spec)
        slot["required_fields"] = references
        slot["referenced_fields"] = references
        slot["template_path"] = f"templates/{filename}"
        slot["sha256"] = file_hashes[filename]
    data["review_note"] = (
        "Пакет содержит работающие draft-only starter-шаблоны с каноническими "
        "полями. Перед пилотом каждая форма должна быть заменена или утверждена "
        "уполномоченным специалистом организации."
    )
    manifest_path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", "utf-8")




def sync_public_assets_and_frontend_catalog() -> None:
    public_root = ROOT / "public" / "starter-packs"
    if public_root.exists():
        shutil.rmtree(public_root)
    public_root.mkdir(parents=True, exist_ok=True)
    frontend_packs: list[dict[str, object]] = []
    labels = {"hr": "Кадры", "legal": "Право", "accounting": "Бухгалтерия"}
    for pack_dir in sorted(PACK_ROOT.glob("tier1-*-ru")):
        manifest = json.loads((pack_dir / "pack.json").read_text("utf-8"))
        target = public_root / pack_dir.name
        shutil.copytree(pack_dir, target)
        templates = []
        for slot in manifest["template_slots"]:
            template_path = Path(slot["template_path"])
            templates.append({
                "documentId": slot["document_id"],
                "label": slot["label"],
                "fileName": template_path.name,
                "url": f"/starter-packs/{pack_dir.name}/{template_path.as_posix()}",
                "sha256": slot["sha256"],
            })
        frontend_packs.append({
            "id": manifest["pack_id"],
            "name": labels[manifest["domain"]],
            "description": manifest["review_note"],
            "usageMode": manifest["usage_mode"],
            "templates": templates,
        })
    data = (
        "export interface StarterTemplateAsset {\n"
        "  documentId: string;\n"
        "  label: string;\n"
        "  fileName: string;\n"
        "  url: string;\n"
        "  sha256: string;\n"
        "}\n\n"
        "export interface StarterPackAsset {\n"
        "  id: string;\n"
        "  name: string;\n"
        "  description: string;\n"
        "  usageMode: 'draft_only';\n"
        "  templates: StarterTemplateAsset[];\n"
        "}\n\n"
        "export const STARTER_PACKS: StarterPackAsset[] = "
        + json.dumps(frontend_packs, ensure_ascii=False, indent=2)
        + " as StarterPackAsset[];\n"
    )
    target = ROOT / "src" / "data" / "starterPacks.ts"
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(data, "utf-8")


def main() -> int:
    hashes: dict[str, dict[str, str]] = {}
    for spec in SPECS:
        path = build(spec)
        hashes.setdefault(spec.pack, {})[spec.filename] = sha256(path)
    for pack, file_hashes in hashes.items():
        update_manifest(pack, file_hashes)
    sync_public_assets_and_frontend_catalog()
    print(json.dumps(hashes, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
